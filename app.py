import dash
import dash_bootstrap_components as dbc
from dash import dcc, html, Input, Output, State
import time
import base64
import io
import dash_auth
import Dashauth
from docx import Document
from helpers import *
from dash import html

privacy_policy_modal = dbc.Modal(
    [
        dbc.ModalHeader(dbc.ModalTitle("AI and Data Privacy Policy for AI Applications")),
        dbc.ModalBody(
            [
                html.P("Effective Date: 17/03/2025", className="text-muted small"),
                html.P("Last Updated: 17/03/2025", className="text-muted small mb-4"),
                html.H5("Introduction", className="mt-4"),
                html.P(
                    [
                        "At Code for Africa (CfA), we are committed to ensuring the responsible and ethical use of",
                        " Artificial Intelligence (AI) and data in all our applications. This AI and Data Privacy Policy",
                        " outlines the principles, practices, and safeguards we implement to protect user privacy,",
                        " ensure transparency, and uphold ethical standards in the development and deployment of AI technologies.",
                    ]
                ),
                html.P(
                    [
                        "This policy applies to all AI systems, tools, and applications developed, maintained, or utilized",
                        " by Code for Africa, including but not limited to machine learning models, natural language",
                        " processing systems, recommendation engines, and automated decision-making systems.",
                    ]
                ),
                html.P(
                    [
                        "For inquiries, please contact us at:",
                        html.Br(),
                        html.Strong("Code for Africa Secretariat"),
                        html.Br(),
                        "112 Loop Street, Cape Town, Western Cape, 8000, South Africa",
                        html.Br(),
                        "South Africa NPO Number: 168-092",
                    ],
                    className="border-start border-primary ps-3 my-4",
                ),
                html.H5("1. Purpose and Scope", className="mt-4"),
                html.P("The purpose of this policy is to:"),
                html.Ul(
                    [
                        html.Li("Protect the privacy and security of personal data processed by AI systems."),
                        html.Li(
                            [
                                "CfA commits to adhering to:",
                                html.Ul(
                                    [
                                        html.Li("GDPR (EU General Data Protection Regulation)."),
                                        html.Li(
                                            [
                                                "Africa-specific laws:",
                                                html.Ul(
                                                    [
                                                        html.Li("Kenya’s Data Protection Act (2019)."),
                                                        html.Li("Nigeria’s NDPR (Nigeria Data Protection Regulation)."),
                                                        html.Li(
                                                            "South Africa’s POPIA (Protection of Personal Information Act)."
                                                        ),
                                                        html.Li("Other regional laws where CfA operates."),
                                                    ]
                                                ),
                                            ]
                                        ),
                                        html.Li(
                                            "Global AI Ethics Frameworks: OECD AI Principles, UNESCO Recommendations on AI Ethics."
                                        ),
                                    ]
                                ),
                            ]
                        ),
                        html.Li("Promote transparency, fairness, and accountability in the use of AI."),
                        html.Li("Minimize risks associated with bias, discrimination, and misuse of AI technologies."),
                    ]
                ),
                html.P("This policy applies to:"),
                html.Ul(
                    [
                        html.Li(
                            "All employees, contractors, and third-party vendors involved in the design, development, deployment, or maintenance of AI systems."
                        ),
                        html.Li(
                            "All users and stakeholders interacting with AI applications developed by Code for Africa."
                        ),
                    ]
                ),
                html.H5("2. Key Principles", className="mt-4"),
                html.H6("2.1. Lawfulness, Transparency, and Accountability", className="mt-3"),
                html.Ul(
                    [
                        html.Li(
                            [
                                html.Strong("Lawfulness: "),
                                "We will only process personal data in accordance with applicable laws and regulations. AI systems will be designed to comply with legal requirements, including obtaining explicit consent where necessary.",
                            ]
                        ),
                        html.Li(
                            [
                                html.Strong("Transparency: "),
                                "Users will be informed about how their data is collected, processed, and used by AI systems. Clear and accessible explanations will be provided regarding the purpose, scope, and impact of AI-driven decisions.",
                            ]
                        ),
                        html.Li(
                            [
                                html.Strong("Accountability: "),
                                "Code for Africa takes full responsibility for the ethical and lawful use of AI. Regular audits and assessments will be conducted to ensure compliance with this policy.",
                            ]
                        ),
                    ]
                ),
                html.H6("2.2. Data Minimization", className="mt-3"),
                html.P(
                    "AI systems will only collect and process data that is strictly necessary for the intended purpose. Excessive or irrelevant data collection will be avoided to minimize privacy risks."
                ),
                html.H6("2.3. Purpose Limitation", className="mt-3"),
                html.P(
                    "Personal data collected for AI applications will only be used for the specific purposes disclosed to users at the time of collection. Any secondary use of data will require additional consent or justification."
                ),
                html.H6("2.4. Accuracy and Quality", className="mt-3"),
                html.P(
                    "We will ensure that data used in AI systems is accurate, up-to-date, and relevant. Mechanisms will be in place to correct inaccuracies and address outdated information promptly."
                ),
                html.H6("2.5. Security and Confidentiality", className="mt-3"),
                html.P(
                    [
                        "Robust technical and organizational measures will be implemented to protect personal data from unauthorized access, loss, alteration, or disclosure.  These include encryption, access controls, and regular security audits."
                    ]
                ),
                html.H6("2.6. Fairness and Non-Discrimination", className="mt-3"),
                html.P(
                    [
                        "AI systems will be designed to avoid bias, discrimination, or unfair treatment of individuals or groups. Regular testing and monitoring will be conducted to identify and mitigate any unintended biases."
                    ]
                ),
                html.H6("2.7. Human Oversight", className="mt-3"),
                html.P(
                    [
                       "AI systems will incorporate mechanisms for human oversight, especially in high-stakes scenarios such as healthcare, finance, or law enforcement. Automated decisions will be subject to review and appeal processes."
                    ]
                ),

                html.H5("3. Data Collection and Processing", className="mt-4"),
                html.H6("3.1. Types of Data Collected", className="mt-3"),
                html.P(
                    [
                       "AI systems may collect and process the following types of data:",
                       html.Ul([
                           html.Li("Personal Identifiable Information (PII): Names, email addresses, phone numbers, etc."),
                           html.Li("Sensitive Data: Health records, financial information, biometric data, etc."),
                           html.Li("Behavioral Data: User interactions, preferences, and usage patterns."),
                           html.Li("Inferred Data: Insights derived from analyzing raw data (e.g., predictions, recommendations).")
                       ])
                    ]
                ),

                html.H6("3.2. Consent and Opt-Out", className="mt-3"),
                html.P(
                    [
                       "Users will be provided with clear options to opt-in or opt-out of data collection and processing activities.",
                       html.Br(),
                       "Explicit consent will be obtained before collecting sensitive data or using data for purposes beyond the original intent."
                    ]
                ),

                html.H6("3.3. Anonymization and Pseudonymization", className="mt-3"),
                html.P("Where possible, personal data will be anonymized or pseudonymized to reduce the risk of re-identification. Aggregated data may be used for analysis without compromising individual privacy."),

                html.H6("3.4. Retention and Deletion", className="mt-3"),
                html.P("Personal data will only be retained for as long as necessary to fulfill the stated purpose. Users will have the right to request the deletion of their data in accordance with applicable laws."),


                html.H5("4. Ethical Use of AI", className="mt-4"),
                html.H6("4.1. Bias Mitigation", className="mt-3"),
                html.P([
                    "AI models will be trained on diverse and representative datasets to minimize bias.",
                    html.Br(),
                    "Regular audits will be conducted to identify and address any discriminatory outcomes."
                ]),
                html.H6("4.2. Explainability", className="mt-3"),
                html.P("AI systems will provide explanations for their decisions and recommendations, especially in critical domains like healthcare, education, and employment. Efforts will be made to ensure that these explanations are understandable to non-technical users."),
                html.H6("4.3. Impact Assessments", className="mt-3"),
                html.P("Before deploying AI systems, comprehensive impact assessments will be conducted to evaluate potential risks to privacy, fairness, and societal well-being."),
                html.H6("4.4. Prohibition of Harmful Use", className="mt-3"),
                 html.P("AI systems will not be used for purposes that could cause harm, violate human rights, or undermine democratic values. Examples include surveillance, manipulation, or exploitation."),

                html.H5("5. Third-Party Vendors and Partners", className="mt-4"),
                html.H6("5.1. Vendor Compliance", className="mt-3"),
                html.P([
                    "Third-party vendors and partners involved in AI development or data processing must adhere to this policy and demonstrate compliance with relevant data protection laws.",
                    html.Br(),
                    "Contracts with vendors will include clauses requiring adherence to privacy and security standards."
                ]),
                html.H6("5.2. Data Sharing", className="mt-3"),
                html.P("Personal data will not be shared with third parties without explicit user consent, except where required by law or for legitimate business purposes (e.g., fraud prevention)."),

                html.H5("6. User Rights", className="mt-4"),
                html.P([
                    "Users have the following rights regarding their personal data:",
                    html.Ul([
                        html.Li("Right to Access: Users can request access to their personal data and obtain information about how it is being processed."),
                        html.Li("Right to Rectification: Users can request corrections to inaccurate or incomplete data."),
                        html.Li("Right to Erasure: Users can request the deletion of their personal data."),
                        html.Li("Right to Object: Users can object to the processing of their data for specific purposes."),
                        html.Li("Right to Data Portability: Users can request a copy of their data in a machine-readable format.")
                    ]),
                    "Requests related to these rights will be addressed within [timeframe, e.g., 30 days] and free of charge, unless the request is excessive or unfounded."  #Added timeframe placeholder
                ]),

                html.H5("7. Security Measures", className="mt-4"),
                html.P([
                    "To safeguard personal data and AI systems, we implement the following measures:",
                    html.Ul([
                        html.Li("Encryption: Data is encrypted during storage and transmission."),
                        html.Li("Access Controls: Only authorized personnel have access to personal data and AI systems."),
                        html.Li("Regular Audits: Security audits and vulnerability assessments are conducted periodically."),
                        html.Li("Incident Response: A robust incident response plan is in place to address data breaches or system failures.")
                    ])
                ]),

                html.H5("8. Training and Awareness", className="mt-4"),
                html.P([
                    "All employees and stakeholders involved in AI development and deployment will receive training on:",
                    html.Ul([
                        html.Li("Data protection laws and regulations."),
                        html.Li("Ethical AI principles and best practices."),
                        html.Li("Privacy-by-design methodologies."),
                        html.Li("Incident reporting and response procedures.")
                    ])
                ]),

                html.H5("9. Monitoring and Compliance", className="mt-4"),
                html.H6("9.1. Regular Audits", className="mt-3"),
                html.P([
                    "Internal and external audits will be conducted to ensure compliance with this policy and applicable laws.",
                    html.Br(),
                    "Findings from audits will be used to improve AI systems and data handling practices."
                ]),

                html.H6("9.2. Reporting Violations", className="mt-3"),
                html.P([
                    "Any violations of this policy must be reported immediately to the designated Data Protection Officer (DPO) or equivalent authority.",
                    html.Br(),
                    "Whistleblower protections will be provided to encourage reporting without fear of retaliation."
                ]),

                html.H5("10. Contact Information", className="mt-4"),
                html.P("For questions, concerns, or requests related to this policy, please contact:"),
                html.P(
                    [
                        html.Strong("Data Protection Officer (DPO):"),
                        html.Br(),
                        "Email: dpo@codeforafrica.org",
                        html.Br(),
                        "Address: Code for Africa Secretariat",
                        html.Br(),
                        "112 Loop Street, Cape Town, Western Cape, 8000, South Africa",
                        html.Br(),
                        "South Africa NPO Number: 168-092",
                    ],
                    className="border-start border-primary ps-3 my-3",
                ),
                html.H5("Conclusion", className="mt-4"),
                html.P(
                    [
                        "Code for Africa is dedicated to leveraging AI responsibly and ethically while prioritizing user privacy and data protection.",
                        " This policy reflects our commitment to transparency, fairness, and accountability in all AI-related activities.",
                        " By adhering to these principles, we aim to build trust with our users and contribute positively to society.",
                    ]
                ),
                html.Hr(),
                html.P(
                    [
                        html.Strong("Acknowledgment:"),
                        html.Br(),
                        "By using our AI applications, you acknowledge that you have read, understood, and agreed to the terms of this AI and Data Privacy Policy.",
                    ],
                    className="small text-muted",
                ),
            ],
            style={"maxHeight": "70vh", "overflowY": "auto"},
        ),
        dbc.ModalFooter(dbc.Button("Close", id="close-privacy-modal", className="ms-auto")),
    ],
    id="privacy-policy-modal",
    size="lg",
    scrollable=True,
)

# Create Footer with Privacy Policy Button
footer = dbc.Row([
    dbc.Col([
        html.Footer([
            html.P("©️ 2025 Toolkit Generator. Powered by CFA Sandbox AI",
                   className="text-center text-muted mt-4"),
            html.Div([
                dbc.Button("Report Issue",
                           color="link",
                           className="text-muted",
                           href="#",
                           id="report-issue-btn"),
                html.Span(" | "),
                dbc.Button("Privacy Policy",
                           color="link",
                           className="text-muted",
                           id="privacy-policy-btn")
            ], className="text-center small")
        ], className="py-4")
    ])
])

# Combine footer and modal
footer_with_modal = html.Div([
    footer,
    privacy_policy_modal
])



# --- Document Generation Function ---
def generate_toolkit_document(query):
    optimisedQuery = determine_ai_toolkit_type(query)
    sections = [
        ("Introduction to AI and Context", introduction_to_ai_and_context),
        ("AI Applications in the Domain", ai_applications_in_domain),  
        ("AI Readiness and Risk Assessment", ai_readiness_and_risk_assessment),
        ("Practical Guidance for AI Use", practical_guidance_for_ai_use),
        ("Ethical and Responsible AI Use", ethical_and_responsible_ai_use),
        ("Challenges and Future Trends", challenges_and_future_trends),
        ("Resources and Training Materials", resources_and_training_materials),
    ]
    doc = Document()
    set_document_styles(doc)
    doc.add_heading("DOMAIN-SPECIFIC AI TOOLKIT", level=0)  
    add_table_of_contents(doc)
    for title, function in sections:
        content = function(optimisedQuery)
        add_section(doc, title, content, level=1)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- Dash App Setup ---
app = dash.Dash(__name__, external_stylesheets=[dbc.themes.BOOTSTRAP], title="Toolkit Generator") 

server = app.server

auth = dash_auth.BasicAuth(
    app,
    Dashauth.VALID_USERNAME_PASSWORD_PAIRS
)

app.layout = html.Div([
    html.H3("TOOLKIT GENERATOR", style={ 
        'textAlign': 'center',
        'color': '#333366',
        'marginBottom': '16px'
    }),
    html.Div([
        html.Label("Describe your domain and AI needs:", style={'fontWeight': 'bold', 'marginBottom': '10px'}),  # Updated label
        dcc.Textarea(
            id='prompt-input',
            placeholder="Describe the domain (e.g., healthcare, education, agriculture) and your specific needs for an AI toolkit...",  # Added placeholder
            rows=5,
            style={'width': '100%', 'borderRadius': '10px', 'padding': '10px'}
        ),
        html.Br(),
        html.Button(
            'Generate Toolkit',
            id='generate-button',
            n_clicks=0,
            style={
                'backgroundColor': '#4CAF50',
                'color': 'white',
                'padding': '10px 20px',
                'border': 'none',
                'borderRadius': '10px',
                'cursor': 'pointer',
                'marginTop': '10px'
            }
        )
    ], style={'margin': '20px'}),

    dcc.Loading(
        id="loading-1",
        type="default",
        children=[
            html.Div(
                [
                    "👋 Thinking! Generating your domain-specific AI toolkit (this will take approximately 3 minutes) 😇",  # Updated message
                    html.Div(
                        " Please Wait", 
                        style={
                            'textAlign': 'center',
                            'margin': '10px',
                            'fontSize': '18px',
                            'color': '#fff',
                            'backgroundColor': '#28a745',
                            'padding': '10px',
                            'borderRadius': '5px',
                            'display': 'inline-block',
                            'cursor': 'pointer'
                        }
                    )
                ],
                style={'textAlign': 'center', 'margin': '10px', 'fontSize': '18px', 'color': '#333'}
            ),
            html.Br(),
            # html.Div(
            #     [
            #         "This AI tool can make mistakes. Please double-check responses",
            #     ],
            #     style={'textAlign': 'center', 'margin': '10px', 'fontSize': '10px', 'color': '#333'}
            # ),
            html.Div(id='progress-container', style={'margin': '20px'}),
            html.Div(id='chain-of-thought-output', style={
                'whiteSpace': 'pre-line',
                'margin': '20px',
                'backgroundColor': '#f4f4f4',
                'padding': '15px',
                'borderRadius': '10px'
            }),
            html.Div(id='download-link-container', style={'margin': '20px'})
        ]
    ),
    html.Div(
                [
                    "This AI tool can make mistakes. Please double-check responses",
                ],
                style={'textAlign': 'center', 'margin': '10px', 'fontSize': '10px', 'color': '#333'}
            ),

    html.Link(
        rel="stylesheet",
        href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0-beta3/css/all.min.css"
    ),

    # Footer
footer_with_modal
], style={
    'fontFamily': 'Arial',
    'width': '75%',
    'margin': '50px auto',
    'padding': '20px',
    'backgroundColor': '#fff',
    'borderRadius': '10px',
    'boxShadow': '0 0 10px rgba(0,0,0,0.1)'
})


@app.callback(
    [Output('progress-container', 'children'),
     Output('chain-of-thought-output', 'children'),
     Output('download-link-container', 'children'),
     Output('generate-button', 'disabled')],
    [Input('generate-button', 'n_clicks')],
    [State('prompt-input', 'value')]
)
def update_output(n_clicks, prompt):
    if n_clicks == 0 or prompt is None:
        return "", "", "", False

    # Simulate a long-running process with progress updates
    progress_updates = []
    thought_process = []

    # Initial Progress
    progress_updates.append(html.Div([
        dbc.Progress(id='progress-bar', value=0, max=100, style={'width': '100%'}),
        html.P("Initializing process...")
    ]))
    thought_process.append("1. Optimizing your query for the AI model...\n")
    time.sleep(2)

    # Stage 1: Optimizing query
    progress_updates.append(html.Div([
        dbc.Progress(id='progress-bar', value=10, max=100, style={'width': '100%'}),
        html.P("Optimizing query...")
    ]))

    # Simulate section generation (stages 2-8)
    section_names = [
        "Introduction to AI and Context", "AI Applications in the Domain", 
        "AI Readiness and Risk Assessment", "Practical Guidance for AI Use",
        "Ethical and Responsible AI Use", "Challenges and Future Trends",
        "Resources and Training Materials"
    ]
    for i, section_name in enumerate(section_names):
        thought_process.append(f"{i + 2}. Generating section: {section_name}...\n")
        time.sleep(2)
        progress_value = 10 + (i + 1) * (70 / len(section_names))
        progress_updates.append(html.Div([
            dbc.Progress(id='progress-bar', value=progress_value, max=100, style={'width': '100%'}),
            html.P(f"Generating {section_name}...")
        ]))

    # Stage 9: Assembling document
    thought_process.append("9. Assembling the document...\n")
    time.sleep(2)
    progress_updates.append(html.Div([
        dbc.Progress(id='progress-bar', value=90, max=100, style={'width': '100%'}),
        html.P("Finalizing document...")
    ]))

    # Stage 10: Preparing document for download
    thought_process.append("10. Preparing document for download...\n")
    buffer = generate_toolkit_document(prompt)
    progress_updates.append(html.Div([
        dbc.Progress(id='progress-bar', value=100, max=100, style={'width': '100%'}),
        html.P("Document ready!")
    ]))

    # Create a base64 encoded string of the document
    buffer.seek(0)
    doc_bytes = buffer.read()
    b64 = base64.b64encode(doc_bytes).decode()
    download_link = html.A(
        'Download Your Domain-Specific AI Toolkit', 
        id='download-link',
        download="Domain_Specific_AI_Toolkit.docx", 
        href=f"data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}",
        target="_blank",
        style={
            'backgroundColor': '#007BFF',
            'color': 'white',
            'padding': '10px 20px',
            'borderRadius': '5px',
            'textDecoration': 'none',
            'display': 'inline-block'
        }
    )

    return progress_updates, "".join(thought_process), download_link, False

@app.callback(
    Output("privacy-policy-modal", "is_open"),
    [Input("privacy-policy-btn", "n_clicks"), Input("close-privacy-modal", "n_clicks")],
    [State("privacy-policy-modal", "is_open")],
)
def toggle_modal(n1, n2, is_open):
    if n1 or n2:
        return not is_open
    return is_open


if __name__ == '__main__':
    app.run_server()
