import time
from docx import Document
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openai import OpenAI
import os
from dotenv import load_dotenv

load_dotenv()

# YOUR_API_KEY = os.getenv("YOUR_API_KEY")
YOUR_API_KEY="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"

def determine_ai_toolkit_type(user_input):
    messages = [
        {
            "role": "system",
            "content": (
                "You are an AI toolkit advisor, specializing in creating domain-specific AI solutions. Your task is to determine the type of AI toolkit a user wants to build based on their input. The final output should categorize their needs and suggest relevant AI applications. Ensure the language follows UK English conventions and that the content is structured, actionable, and engaging."
            ),
        },
        {
            "role": "user",
            "content": (
                f"Input: {user_input}\n\n"
                "Task: Analyse the user's input and determine the appropriate AI toolkit category. Provide recommendations in the following structure:\n\n"
                "1. Identified AI Toolkit Category\n"
                "   - Classify the toolkit type (e.g., AI for Education, AI for Business Analytics, AI for Healthcare, AI for Agriculture, etc.).\n\n"
                "2. Core AI Technologies Required\n"
                "   - List essential AI technologies such as Machine Learning, Natural Language Processing, Computer Vision, etc.\n\n"
                "3. Relevant AI Use Cases\n"
                "   - Describe key applications for the selected domain (e.g., personalized learning for education, predictive analytics for business, diagnostic support for healthcare).\n\n"
                "4. Recommended Tools and Platforms\n"
                "   - Provide an overview of suitable AI tools, including open-source and commercial options.\n\n"
                "5. Implementation Considerations\n"
                "   - Highlight challenges, best practices, and ethical considerations.\n\n"
                "THE OUTPUT SHOULD BE CLEAN TEXT WITHOUT '#' "
                "Output: An 8 sentence summary of the above points"
            ),
        }
    ]

    client = OpenAI(api_key=YOUR_API_KEY, base_url="https://api.perplexity.ai")

    response = client.chat.completions.create(
        model="sonar-pro",
        messages=messages,
        # timeout=60
    )

    extracted_data = response.choices[0].message.content

    return extracted_data



def introduction_to_ai_and_context(query):
    messages = [
        {
            "role": "system",
            "content": (
                "You are a chief researcher and an expert in building AI-Focused Toolkits that help professionals improve productivity "
                "and integrate open-source tools into their workflows. Your task is to write an introduction to AI and its context in the specific domain, "
                "tailored for professionals in that field. "
                "The final document must be in UK English, not exceed 3 pages, use Times Roman font, "
                "and cite sources via clickable links rather than inline referencing. Tailor the content to be engaging and actionable for the target audience."
            ),
        },
        {
            "role": "user",
            "content": (
                f"Input: {query}\n\n"
                "Task: Write an introduction to AI and its context in the specified domain using the following guidelines:\n\n"
                "1.  Key AI Concepts\n"
                "    -   Core AI Technologies: Artificial Intelligence, Machine Learning, Generative AI, Natural Language Processing (NLP).\n"
                "    -   Discuss the above terms using at least 3 lines each and make it easy to understand for professionals in the domain.\n"
                "    -   For each term provide at least 2 citations/urls.\n"
                "    -   AI's Role in the Domain: How AI transforms tasks, improves accuracy, and drives efficiency.\n"
                "        * Provide contextual examples relevant to the domain (e.g., use of AI for specific domain challenges).\n"
                "2.  Domain-Specific Concepts\n"
                "    -   Relevant Jargon and Terms: Introduce specialised terminology, explain their meanings, and show how they relate to AI tools.\n"
                "    -   Examples and Use Cases: Provide context for how these terms apply in real-world scenarios.\n"
                "        * Focus on challenges and opportunities specific to the domain (e.g., addressing domain-specific problems, enhancing processes).\n"
                "THE OUTPUT SHOULD BE CLEAN TEXT WITHOUT '#' "
            ),
        }
    ]

    client = OpenAI(api_key=YOUR_API_KEY, base_url="https://api.perplexity.ai")
    response = client.chat.completions.create(
        model="sonar-pro",
        messages=messages,
        # timeout=60
    )

    extracted_data = response.choices[0].message.content

    return extracted_data


def ai_applications_in_domain(query):
    messages = [
        {
            "role": "system",
            "content": (
                "You are a chief researcher and an expert in building AI-Focused Toolkits that help professionals improve productivity "
                "and integrate open-source tools into their workflows. Your task is to write a section on AI applications in the specified domain, "
                "with relevant context and examples. "
                "The final document must be in UK English, not exceed 4 pages, use Times Roman font, "
                "and cite sources via clickable links rather than inline referencing. Tailor the content to be engaging and actionable for the target audience."
            ),
        },
        {
            "role": "user",
            "content": (
                f"Input: {query}\n\n"
                "Task: Write a section on AI applications in the specified domain using the following guidelines:\n\n"
                "1.  Use Cases and Strategies\n"
                "    -   Practical Applications: Showcase practical AI applications in the domain (e.g., automation, data analysis, pattern recognition).\n"
                "        * Provide domain-specific examples that demonstrate AI's impact and potential.\n"
                "    -   AI's Potential: Discuss AI's potential for enhancing productivity, innovation, and outcomes in the domain.\n"
                "2.  Tools and Platforms\n"
                "    -   Overview of Tools: Provide an overview of relevant AI tools for professionals in the domain (cloud-based, open-source, or self-hosted).\n"
                "        * Include at least 32 AI tools across 8 categories that are relevant to the domain.\n"
                "        * Prioritize tools that are easy to use (provide a description), open-source, or cost-effective.\n"
                "THE OUTPUT SHOULD BE CLEAN TEXT WITHOUT '#' "
            ),
        }
    ]

    client = OpenAI(api_key=YOUR_API_KEY, base_url="https://api.perplexity.ai")
    response = client.chat.completions.create(
        model="sonar-pro",
        messages=messages,
        # timeout=60
    )

    extracted_data = response.choices[0].message.content

    return extracted_data


def ai_readiness_and_risk_assessment(query):
    messages = [
        {
            "role": "system",
            "content": (
                "You are a chief researcher and an expert in building AI-Focused Toolkits that help professionals improve productivity "
                "and integrate open-source tools into their workflows. Your task is to write a section on AI readiness and risk assessment for organizations in the specified domain. "
                "The final document must be in UK English, not exceed 3 pages, use Times Roman font, "
                "and cite sources via clickable links rather than inline referencing. Tailor the content to be engaging and actionable for the target audience."
            ),
        },
        {
            "role": "user",
            "content": (
                f"Input: {query}\n\n"
                "Task: Write a section on AI readiness and risk assessment using the following guidelines:\n\n"
                "1.  Adoption Framework\n"
                "    -   Readiness Tests: Provide at least 3 AI readiness tests/guides to assess organisational readiness for AI implementation in the domain.\n"
                "        * Include practical evaluation frameworks, checklists, and self-assessment tools with scoring systems.\n"
                "        * Offer tests that cover technical infrastructure, data capabilities, staff skills, and organizational culture.\n"
                "        * Provide detailed methodology for each test, including how to interpret results.\n"
                "    -   Risk Assessment: Provide at least 3 risk assessment tests/guides to identify potential pitfalls and ethical concerns specific to the domain context.\n"
                "        * Include comprehensive risk matrices with impact and probability ratings.\n"
                "        * Offer domain-specific risk identification questionnaires with actionable mitigation strategies.\n"
                "        * Provide benchmarking tools to compare risk levels with industry standards.\n"
                "2.  Evaluation Metrics\n"
                "    -   Metrics for Success: Define metrics for evaluating the success of AI adoption in the domain (e.g., increased efficiency, improved accuracy, enhanced processes).\n"
                "THE OUTPUT SHOULD BE CLEAN TEXT WITHOUT '#' "
            ),
        }
    ]

    client = OpenAI(api_key=YOUR_API_KEY, base_url="https://api.perplexity.ai")
    response = client.chat.completions.create(
        model="sonar-pro",
        messages=messages,
        # timeout=60
    )

    extracted_data = response.choices[0].message.content

    return extracted_data


def practical_guidance_for_ai_use(query):
    messages = [
        {
            "role": "system",
            "content": (
                "You are a chief researcher and an expert in building AI-Focused Toolkits that help professionals improve productivity "
                "and integrate open-source tools into their workflows. Your task is to write practical guidance on using AI tools in the specified domain, "
                "including prompt engineering and operational considerations. "
                "The final document must be in UK English, not exceed 3 pages, use Times Roman font, "
                "and cite sources via clickable links rather than inline referencing. Tailor the content to be engaging and actionable for the target audience."
            ),
        },
        {
            "role": "user",
            "content": (
                f"Input: {query}\n\n"
                "Task: Write practical guidance on AI use in the specified domain using the following guidelines:\n\n"
                "1.  Prompt Engineering\n"
                "    -   Effective Prompts: Offer guidance on crafting effective prompts for large language models (LLMs) and AI systems, including templates and engaging examples relevant to the domain.\n"
                "2.  Operational Considerations\n"
                "    -   Workflow Integration: Explain how to integrate AI tools into existing workflows in organizations within the domain.\n"
                "    -   Challenges: Address challenges such as AI bias, data quality, and scalability in the domain context.\n"
                "3.  Examples of Resources and Databases\n"
                "    -   Opensource materials for professionals in the domain.\n"
                "THE OUTPUT SHOULD BE CLEAN TEXT WITHOUT '#' "
            ),
        }
    ]

    client = OpenAI(api_key=YOUR_API_KEY, base_url="https://api.perplexity.ai")
    response = client.chat.completions.create(
        model="sonar-pro",
        messages=messages,
        # timeout=60
    )

    extracted_data = response.choices[0].message.content

    return extracted_data


def ethical_and_responsible_ai_use(query):
    messages = [
        {
            "role": "system",
            "content": (
                "You are a chief researcher and an expert in building AI-Focused Toolkits that help professionals improve productivity "
                "and integrate open-source tools into their workflows. Your task is to write a section on ethical and responsible AI use in the specified domain. "
                "The final document must be in UK English, not exceed 3 pages, use Times Roman font, "
                "and cite sources via clickable links rather than inline referencing. Tailor the content to be engaging and actionable for the target audience."
            ),
        },
        {
            "role": "user",
            "content": (
                f"Input: {query}\n\n"
                "Task: Write a section on ethical and responsible AI use in the specified domain using the following guidelines:\n\n"
                "1.  Ethical Principles\n"
                "    -   Core Values: Discuss fairness, transparency, and accountability in AI deployment within the domain context.\n"
                "    -   Critical Issues: Address verification, misinformation, bias mitigation, and other ethical concerns specific to the domain.\n"
                "2.  Regulatory Considerations\n"
                "    -   Compliance: Highlight compliance with relevant standards and legal frameworks applicable to the domain.\n"
                "THE OUTPUT SHOULD BE CLEAN TEXT WITHOUT '#' "
            ),
        }
    ]

    client = OpenAI(api_key=YOUR_API_KEY, base_url="https://api.perplexity.ai")
    response = client.chat.completions.create(
        model="sonar-pro",
        messages=messages,
        # timeout=60
    )

    extracted_data = response.choices[0].message.content

    return extracted_data


def challenges_and_future_trends(query):
    messages = [
        {
            "role": "system",
            "content": (
                "You are a chief researcher and an expert in building AI-Focused Toolkits that help professionals improve productivity "
                "and integrate open-source tools into their workflows. Your task is to write a section on challenges and future trends of AI in the specified domain. "
                "The final document must be in UK English, not exceed 3 pages, use Times Roman font, "
                "and cite sources via clickable links rather than inline referencing. Tailor the content to be engaging and actionable for the target audience."
            ),
        },
        {
            "role": "user",
            "content": (
                f"Input: {query}\n\n"
                "Task: Write a section on challenges and future trends of AI in the specified domain using the following guidelines:\n\n"
                "1.  Current Challenges\n"
                "    -   Limitations: Discuss limitations of AI tools in the domain (e.g., accuracy, scalability, over-reliance).\n"
                "        * Focus on challenges specific to the domain (e.g., data limitations, integration issues, skill gaps).\n"
                "2.  Emerging Trends\n"
                "    -   Future Advancements: Explore future possibilities, such as advancements in AI models or new applications relevant to the domain.\n"
                "3.  Case Studies\n"
                "    -   Examples: Include case studies showing both successful and failed implementations of AI in the domain.\n"
                "THE OUTPUT SHOULD BE CLEAN TEXT WITHOUT '#' "
            ),
        }
    ]

    client = OpenAI(api_key=YOUR_API_KEY, base_url="https://api.perplexity.ai")
    response = client.chat.completions.create(
        model="sonar-pro",
        messages=messages,
        # timeout=60
    )

    extracted_data = response.choices[0].message.content

    return extracted_data


def resources_and_training_materials(query):
    messages = [
        {
            "role": "system",
            "content": (
                "You are a chief researcher and an expert in building AI-Focused Toolkits that help professionals improve productivity "
                "and integrate open-source tools into their workflows. Your task is to write a section on resources and training materials for professionals interested in AI in the specified domain. "
                "The final document must be in UK English, not exceed 3 pages, use Times Roman font, "
                "and cite sources via clickable links rather than inline referencing. Tailor the content to be engaging and actionable for the target audience."
            ),
        },
        {
            "role": "user",
            "content": (
                f"Input: {query}\n\n"
                "Task: Write a section on resources and training materials using the following guidelines:\n\n"
                "1.  Further Learning\n"
                "    -   Provide 10 curated further learning materials (links to research papers, articles, courses, etc.) relevant to AI in the domain.\n"
                "2.  Exercises and Workshops\n"
                "    -   Include 5 hands-on exercises or workshop guides for practical learning about AI tools and techniques, tailored for professionals in the domain.\n"
                "3.  Case Studies\n"
                "    -   Present 5 real-world case studies illustrating effective AI deployment in the domain from various regions and contexts.\n"
                "THE OUTPUT SHOULD BE CLEAN TEXT WITHOUT '#' "
            ),
        }
    ]

    client = OpenAI(api_key=YOUR_API_KEY, base_url="https://api.perplexity.ai")
    response = client.chat.completions.create(
        model="sonar-pro",
        messages=messages,
        # timeout=60
    )

    extracted_data = response.choices[0].message.content

    return extracted_data


def clean_and_format_document(document_content):
    messages = [
        {
            "role": "system",
            "content": (
                "You are a professional document formatter and editor specializing in technical and academic documents. "
                "Your expertise is in transforming raw document content into clean, well-structured, and professionally formatted documents. "
                "You maintain all original information while significantly improving readability, organization, and visual appeal."
            ),
        },
        {
            "role": "user",
            "content": (
                f"Input: {document_content}\n\n"
                "Please format and structure this Domain-Specific AI Toolkit document to create a clean, professional, and readable resource. Follow these guidelines:\n\n"
                "1. Create a consistent heading hierarchy using proper Markdown formatting (# for main headings, ## for subheadings, etc.)\n"
                "2. Format all lists consistently, using either numbered or bulleted lists as appropriate for the content\n"
                "3. Standardize the formatting for:\n"
                "   - Citations and references\n"
                "   - Case studies\n"
                "   - Examples and code snippets\n"
                "   - Tool descriptions\n"
                "4. Create a clean table of contents with hyperlinks to each section\n"
                "5. For each major section:\n"
                "   - Add a brief introduction that summarizes key points\n"
                "   - Ensure proper spacing between paragraphs and sections\n"
                "   - Format any tables using proper Markdown table syntax\n"
                "6. Highlight important concepts, tools, or metrics using bold or italic formatting where appropriate\n"
                "7. Remove any redundant text, inconsistent formatting, or stray characters (like \\*)\n"
                "8. Ensure all bullet points and numbered lists align properly\n"
                "9. Add proper spacing before and after headings, lists, and code blocks\n"
                "10. Create a consistent style for callouts or special notes\n\n"
                "Please preserve all the original content and information while making these formatting improvements."
            ),
        }
    ]

    client = OpenAI(api_key=YOUR_API_KEY, base_url="https://api.perplexity.ai")
    response = client.chat.completions.create(
        model="sonar-pro",
        messages=messages,
    )

    formatted_document = response.choices[0].message.content

    return formatted_document


# Creat & Style docx!

def set_document_styles(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_table_of_contents(doc):
    doc.add_paragraph("Table of Contents", style='Heading 1')
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    run = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    doc.add_page_break()


def add_section(doc, title, content, level=1):
    doc.add_heading(title, level=level)
    doc.add_paragraph(content)
    doc.add_paragraph()

if __name__ == "__main__":
    pass
